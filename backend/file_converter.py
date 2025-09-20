import os
import tempfile
import subprocess
import logging
from pathlib import Path
from typing import Optional
import pypandoc
from docx import Document
import PyPDF2
from io import BytesIO
from pdf2docx import Converter

logger = logging.getLogger(__name__)

class FileConverter:
    """Handle file format conversions"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
    
    async def convert_file(self, input_path: str, input_format: str, output_format: str, conversion_id: str) -> str:
        """Convert file from input format to output format"""
        try:
            output_filename = f"{conversion_id}_converted.{output_format}"
            output_path = os.path.join(self.temp_dir, output_filename)
            
            # Route to appropriate conversion method
            if input_format == "pdf":
                if output_format == "docx":
                    await self._convert_pdf_to_docx_professional(input_path, output_path)
                elif output_format in ["txt", "html"]:
                    await self._convert_pdf_to_text_based(input_path, output_path, output_format)
                else:
                    await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
            
            elif input_format in ["docx", "doc", "rtf", "odt"]:
                if output_format == "pdf":
                    await self._convert_to_pdf_with_libreoffice(input_path, output_path)
                elif output_format in ["txt", "html", "rtf", "odt", "docx"]:
                    await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
                else:
                    await self._copy_file(input_path, output_path)
            
            elif input_format == "txt":
                await self._convert_text_based(input_path, output_path, input_format, output_format)
            
            else:
                # Use pandoc for other formats
                await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
            
            if not os.path.exists(output_path):
                raise Exception(f"Conversion failed: Output file not created")
            
            logger.info(f"Successfully converted {input_format} to {output_format}")
            return output_path
            
        except Exception as e:
            logger.error(f"Conversion error: {str(e)}")
            raise Exception(f"Failed to convert file: {str(e)}")

    async def _convert_to_pdf_with_libreoffice(self, input_path: str, output_path: str):
        """Convert document to PDF using LibreOffice for better fidelity"""
        try:
            # The output directory for LibreOffice must exist
            output_dir = os.path.dirname(output_path)

            cmd = [
                "libreoffice",
                "--headless",
                "--convert-to", "pdf",
                "--outdir", output_dir,
                input_path
            ]
            subprocess.run(cmd, check=True, capture_output=True, text=True)

            # Rename the output file to the expected name
            base_filename = Path(input_path).stem
            expected_pdf_path = os.path.join(output_dir, f"{base_filename}.pdf")
            if os.path.exists(expected_pdf_path):
                os.rename(expected_pdf_path, output_path)
            else:
                raise Exception("LibreOffice conversion did not produce the expected PDF file.")

        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: {e.stderr}")
            # Fallback to pandoc if LibreOffice fails
            await self._convert_with_pandoc(input_path, output_path, Path(input_path).suffix[1:], "pdf")
        except Exception as e:
            logger.error(f"An unexpected error occurred during LibreOffice conversion: {e}")
            await self._convert_with_pandoc(input_path, output_path, Path(input_path).suffix[1:], "pdf")

    async def _convert_pdf_to_docx_professional(self, input_path: str, output_path: str):
        """Convert PDF to DOCX using pdf2docx for better formatting retention"""
        try:
            cv = Converter(input_path)
            cv.convert(output_path, start=0, end=None)
            cv.close()
        except Exception as e:
            logger.error(f"pdf2docx conversion failed: {str(e)}")
            # Fallback to simple text extraction if professional conversion fails
            await self._convert_pdf_to_docx_simple(input_path, output_path)

    async def _convert_with_pandoc(self, input_path: str, output_path: str, input_format: str, output_format: str):
        """Convert using pandoc"""
        try:
            format_mapping = {
                "docx": "docx", "doc": "doc", "txt": "plain", "html": "html",
                "rtf": "rtf", "odt": "odt", "pdf": "pdf"
            }
            input_fmt = format_mapping.get(input_format, input_format)
            output_fmt = format_mapping.get(output_format, output_format)
            
            pypandoc.convert_file(
                input_path, output_fmt, outputfile=output_path, format=input_fmt
            )
        except Exception as e:
            try:
                cmd = ["pandoc", input_path, "-f", input_format, "-t", output_format, "-o", output_path]
                subprocess.run(cmd, check=True, capture_output=True, text=True)
            except subprocess.CalledProcessError as e:
                raise Exception(f"Pandoc conversion failed: {e.stderr}")

    async def _convert_pdf_to_text_based(self, input_path: str, output_path: str, output_format: str):
        """Convert PDF to text-based formats"""
        try:
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
            
            if output_format == "txt":
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(text)
            elif output_format == "html":
                html_content = f"<!DOCTYPE html><html><head><title>Converted Document</title><meta charset=\"UTF-8\"></head><body><pre>{text}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
        except Exception as e:
            raise Exception(f"PDF to text conversion failed: {str(e)}")

    async def _convert_pdf_to_docx_simple(self, input_path: str, output_path: str):
        """Convert PDF to DOCX using simple text extraction"""
        try:
            with open(input_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = "".join(page.extract_text() + "\n" for page in pdf_reader.pages)
            
            doc = Document()
            doc.add_paragraph(text)
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"PDF to DOCX (simple) conversion failed: {str(e)}")

    async def _copy_file(self, input_path: str, output_path: str):
        """Copy a file for same-format 'conversion'"""
        import shutil
        shutil.copy2(input_path, output_path)

    async def _convert_text_based(self, input_path: str, output_path: str, input_format: str, output_format: str):
        """Convert between text-based formats"""
        try:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if output_format == "html":
                html_content = f"<!DOCTYPE html><html><head><title>Converted Document</title><meta charset=\"UTF-8\"></head><body><pre>{content}</pre></body></html>"
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
            elif output_format == "docx":
                doc = Document()
                doc.add_paragraph(content)
                doc.save(output_path)
            else:
                await self._convert_with_pandoc(input_path, output_path, input_format, output_format)
        except Exception as e:
            raise Exception(f"Text conversion failed: {str(e)}")